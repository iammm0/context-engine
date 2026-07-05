"""Word文档解析器"""
from typing import Dict, Any, List
import os
import re
import logging
from .base import BaseParser

logger = logging.getLogger(__name__)

# 检查可用的库
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def _build_word_image_ocr_metadata(images_info: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Build normalized OCR metadata for embedded Word images."""
    ocr_text_parts = [
        str(image.get("ocr_text") or "").strip()
        for image in images_info
        if str(image.get("ocr_text") or "").strip()
    ]
    return {
        "image_count": len(images_info),
        "ocr_text_length": len("\n\n".join(ocr_text_parts)),
        "images": [
            {
                "image_index": image.get("image_index"),
                "target": image.get("target"),
                "confidence": image.get("confidence"),
                "line_count": image.get("line_count"),
                "text_length": len(str(image.get("ocr_text") or "")),
                "width": image.get("width"),
                "height": image.get("height"),
            }
            for image in images_info
        ],
    }


class WordParser(BaseParser):
    """Word文档解析器（支持.docx和.doc格式）"""
    
    def _clean_text(self, text: str) -> str:
        """
        清理文本，处理特殊字符和公式，避免乱码
        优化：保留数学公式的LaTeX格式
        
        Args:
            text: 原始文本
            
        Returns:
            清理后的文本
        """
        if not text:
            return ""
        
        # 确保是字符串类型
        if isinstance(text, bytes):
            try:
                text = text.decode('utf-8', errors='replace')
            except Exception:
                try:
                    text = text.decode('gbk', errors='replace')  # 尝试GBK编码（中文常用）
                except Exception:
                    text = text.decode('latin-1', errors='replace')
        
        # 先提取并保护公式
        from utils.formula_extractor import FormulaExtractor
        text = FormulaExtractor.preserve_formulas_in_text(text)
        
        # 统一换行符
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # 移除控制字符（保留换行、制表符和常见空白字符）
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # 移除 Word 文档中的嵌入对象标记（如 EMBEDEquation, EMBED 等）
        # 但保留已标记的公式
        # 只在非公式区域移除这些标记
        parts = re.split(r'(\$\$.*?\$\$|\$[^\$]+\$)', text)
        cleaned_parts = []
        for part in parts:
            if part.startswith('$') and part.endswith('$'):
                cleaned_parts.append(part)  # 保留公式
            else:
                # 移除嵌入对象标记
                part = re.sub(r'EMBED\w*\.?\d*', '', part, flags=re.IGNORECASE)
                part = re.sub(r'EMBEDEquation[\.\s]*\d*', '', part, flags=re.IGNORECASE)
                # 移除常见的格式标记和控制序列
                part = re.sub(r'[!\'"][a-z]{1,3}[a-z]{1,3}\s*\d+', '', part, flags=re.IGNORECASE)
                part = re.sub(r'\d{4,}[A-Z]{4,}', '', part)
                cleaned_parts.append(part)
        text = ''.join(cleaned_parts)
        
        # 处理常见的编码问题
        replacements = {
            '\ufffd': '',  # Unicode替换字符
            '\u200b': '',  # 零宽空格
            '\u200c': '',  # 零宽非连字符
            '\u200d': '',  # 零宽连字符
            '\ufeff': '',  # 字节顺序标记
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        # 移除明显的二进制数据残留（连续的特殊字符，但不是中文标点）
        # 但保留公式标记
        parts = re.split(r'(\$\$.*?\$\$|\$[^\$]+\$)', text)
        cleaned_parts = []
        for part in parts:
            if part.startswith('$') and part.endswith('$'):
                cleaned_parts.append(part)  # 保留公式
            else:
                # 移除二进制残留，但保留常见标点
                part = re.sub(r'[^\w\s\u4e00-\u9fff\u3000-\u303f\uff00-\uffef.,;:!?()\[\]{}\-+=*/<>""''（）【】《》]{3,}', ' ', part)
                cleaned_parts.append(part)
        text = ''.join(cleaned_parts)
        
        # 规范化空白字符
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            if '$$' in line or '$' in line:
                # 保护公式部分
                parts = re.split(r'(\$\$.*?\$\$|\$[^\$]+\$)', line)
                cleaned_parts = []
                for part in parts:
                    if part.startswith('$') and part.endswith('$'):
                        cleaned_parts.append(part)
                    else:
                        cleaned_parts.append(re.sub(r'[ \t]+', ' ', part))
                cleaned_line = ''.join(cleaned_parts)
            else:
                cleaned_line = re.sub(r'[ \t]+', ' ', line)
            cleaned_lines.append(cleaned_line)
        text = '\n'.join(cleaned_lines)
        
        # 移除行首行尾的空白（但保留公式标记）
        lines = text.split('\n')
        final_lines = []
        for line in lines:
            if '$$' in line or '$' in line:
                final_lines.append(line)  # 保留包含公式的行
            else:
                line = line.strip()
                # 跳过只包含特殊字符或数字的行（可能是格式残留）
                if line and len(line) > 1:
                    if re.search(r'[a-zA-Z\u4e00-\u9fff]', line):
                        final_lines.append(line)
        
        return '\n'.join(final_lines)
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析Word文档"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # 根据文件扩展名选择解析方法
        if file_ext == ".docx":
            return self._parse_docx(file_path)
        elif file_ext == ".doc":
            return self._parse_doc(file_path)
        else:
            raise ValueError(f"不支持的文件扩展名: {file_ext}")
    
    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """解析.docx格式"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Install it with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            # 运行时开关（同步解析路径）
            try:
                from services.runtime_config import get_runtime_config_sync

                _cfg = get_runtime_config_sync()
                _modules = _cfg.get("modules") or {}
                ocr_enabled = bool(_modules.get("ocr_image_enabled", True))
                table_enabled = bool(_modules.get("table_parse_enabled", True))
            except Exception:
                ocr_enabled = True
                table_enabled = True
            
            # 提取文本内容，确保使用UTF-8编码
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # 清理文本，处理特殊字符和公式
                    text = self._clean_text(text)
                    if text.strip():
                        paragraphs.append(text)
            
            full_text = "\n\n".join(paragraphs)
            
            # 提取元数据
            core_props = doc.core_properties
            metadata = {
                "title": (core_props.title or "").encode('utf-8', errors='replace').decode('utf-8'),
                "author": (core_props.author or "").encode('utf-8', errors='replace').decode('utf-8'),
                "subject": (core_props.subject or "").encode('utf-8', errors='replace').decode('utf-8'),
                "paragraphs": len(paragraphs),
                "extraction_method": "python-docx"
            }
            
            # 增强功能：提取表格（保留格式）
            tables_info = []
            tables_text = []
            if table_enabled:
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = [
                            cell.text.strip().encode('utf-8', errors='replace').decode('utf-8')
                            for cell in row.cells
                        ]
                        table_data.append(row_data)
                        row_text = " | ".join(row_data)
                        tables_text.append(row_text)
                    
                    if table_data:
                        # 提取表格的HTML和Markdown格式
                        try:
                            from utils.table_extractor import TableExtractor
                            semantic_info = TableExtractor.extract_semantic_structure(table_data)
                            html_format = TableExtractor._to_html(table_data)
                            markdown_format = TableExtractor._to_markdown(table_data)
                            
                            tables_info.append({
                                "html": html_format,
                                "markdown": markdown_format,
                                "semantic": semantic_info
                            })
                        except Exception as e:
                            logger.warning(f"表格格式化失败: {e}")
            
            if tables_text:
                full_text += "\n\n" + "\n\n".join(tables_text)
            
            if tables_info:
                metadata["tables"] = tables_info
            
            # 确保最终文本是UTF-8编码
            if isinstance(full_text, bytes):
                full_text = full_text.decode('utf-8', errors='replace')
            
            # 增强功能：分析公式
            formulas_info = []
            try:
                from utils.formula_analyzer import FormulaAnalyzer
                formulas_info = FormulaAnalyzer.extract_all_formulas_info(full_text)
                if formulas_info:
                    metadata["formulas"] = formulas_info
            except Exception as e:
                logger.warning(f"公式分析失败: {e}")
            
            # 增强功能：提取图片（Word文档中的图片）并对内嵌图做 OCR
            images_info = []
            try:
                import tempfile
                for rel in doc.part.rels.values():
                    if "image" not in rel.target_ref:
                        continue
                    image_index = len(images_info) + 1
                    images_info.append({
                        "image_index": image_index,
                        "target": rel.target_ref,
                        "type": "embedded_image"
                    })
                    if ocr_enabled:
                        # 内嵌图 OCR：获取 blob 写入临时文件后调用 image_ocr
                        try:
                            target_part = rel.target_part
                            if target_part is None:
                                continue
                            type_name = type(target_part).__name__
                            if "ImagePart" not in type_name and "Image" not in type_name:
                                continue
                            blob = getattr(target_part, "_blob", None) or getattr(target_part, "blob", None)
                            if not blob:
                                continue
                            # 扩展名：从 target_ref 如 "media/image1.png" 或 partname 取
                            ref = getattr(rel, "target_ref", "") or ""
                            ext = os.path.splitext(ref)[1].lower() if ref else ".png"
                            if ext not in (".jpg", ".jpeg", ".png", ".bmp", ".webp", ".tiff", ".tif", ".gif"):
                                ext = ".png"
                            images_info[-1]["extension"] = ext.lstrip(".")
                            source_image = getattr(target_part, "image", None)
                            if source_image is not None:
                                images_info[-1]["width"] = getattr(source_image, "px_width", None)
                                images_info[-1]["height"] = getattr(source_image, "px_height", None)
                            fd, tmp_path = tempfile.mkstemp(suffix=ext)
                            try:
                                with os.fdopen(fd, "wb") as f:
                                    f.write(blob)
                                from utils.image_ocr import image_ocr
                                ocr_result = image_ocr.extract_text_from_image(tmp_path)
                                ocr_text = (ocr_result.get("text") or "").strip()
                                if ocr_text:
                                    full_text += f"\n\n[图片文字 image={image_index}]\n" + ocr_text
                                if images_info:
                                    images_info[-1]["ocr_text"] = ocr_text
                                    images_info[-1]["confidence"] = ocr_result.get("confidence", 0.0)
                                    images_info[-1]["line_count"] = ocr_result.get("line_count", 0)
                            finally:
                                try:
                                    os.unlink(tmp_path)
                                except Exception:
                                    pass
                        except Exception as img_e:
                            logger.warning(f"Word 内嵌图 OCR 失败: {rel.target_ref}, {img_e}")
                if images_info:
                    metadata["images"] = images_info
                    metadata["image_ocr"] = _build_word_image_ocr_metadata(images_info)
            except Exception as e:
                logger.warning(f"图片提取失败: {e}")
            
            return {
                "text": full_text,
                "metadata": {
                    **metadata,
                    "tables_count": len(doc.tables)
                }
            }
        except Exception as e:
            raise Exception(f"Failed to parse .docx file: {e}")
    
    def _parse_doc(self, file_path: str) -> Dict[str, Any]:
        """
        解析.doc格式（旧版Word格式）
        
        使用系统工具进行转换：
        - Linux/Mac: antiword 或 LibreOffice
        - Windows: LibreOffice
        """
        import subprocess
        import tempfile
        
        # 尝试使用系统工具转换.doc文件
        # 方法1: 尝试使用antiword（Linux/Mac）
        # 方法2: 尝试使用LibreOffice（跨平台）
        
        full_text = None
        extraction_method = None
        
        # 方法1: 尝试使用antiword
        try:
            result = subprocess.run(
                ['antiword', file_path],
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='replace',
                timeout=60
            )
            if result.returncode == 0 and result.stdout:
                full_text = result.stdout
                extraction_method = "antiword"
        except (FileNotFoundError, subprocess.TimeoutExpired):
            # antiword不可用，继续尝试其他方法
            pass
        except Exception as e:
            # 其他错误，记录但继续尝试
            import logging
            logger = logging.getLogger(__name__)
            logger.warning(f"使用antiword解析.doc文件失败: {str(e)}")
        
        # 方法2: 如果antiword失败，尝试使用LibreOffice
        if not full_text:
            try:
                # 创建临时目录用于输出
                with tempfile.TemporaryDirectory() as temp_dir:
                    output_file = os.path.join(temp_dir, "output.txt")
                    
                    # 使用LibreOffice命令行转换.doc为文本
                    # soffice --headless --convert-to txt --outdir <output_dir> <input_file>
                    result = subprocess.run(
                        ['soffice', '--headless', '--convert-to', 'txt', '--outdir', temp_dir, file_path],
                        capture_output=True,
                        text=True,
                        timeout=120
                    )
                    
                    if result.returncode == 0:
                        # 读取转换后的文本文件
                        base_name = os.path.splitext(os.path.basename(file_path))[0]
                        txt_file = os.path.join(temp_dir, f"{base_name}.txt")
                        
                        if os.path.exists(txt_file):
                            with open(txt_file, 'r', encoding='utf-8', errors='replace') as f:
                                full_text = f.read()
                                extraction_method = "libreoffice"
            except (FileNotFoundError, subprocess.TimeoutExpired):
                # LibreOffice不可用
                pass
            except Exception as e:
                import logging
                logger = logging.getLogger(__name__)
                logger.warning(f"使用LibreOffice解析.doc文件失败: {str(e)}")
        
        # 如果所有方法都失败，抛出异常
        if not full_text:
            raise Exception(
                "无法解析.doc文件。请安装以下工具之一：\n"
                "- Linux: sudo apt-get install antiword 或 sudo apt-get install libreoffice\n"
                "- macOS: brew install antiword 或 brew install --cask libreoffice\n"
                "- Windows: 安装LibreOffice并确保soffice在PATH中\n"
                "或者将.doc文件转换为.docx格式后上传。"
            )
        
        # 清理文本，处理特殊字符和公式
        full_text = self._clean_text(full_text)
        
        # 提取基本元数据
        metadata = {
            "extraction_method": extraction_method,
            "file_format": "doc"
        }
        
        # 尝试从文件名提取标题
        file_name = os.path.basename(file_path)
        if file_name:
            metadata["title"] = os.path.splitext(file_name)[0]
        
        return {
            "text": full_text,
            "metadata": metadata
        }
    
    def supported_extensions(self) -> List[str]:
        return ["docx", "doc"]

